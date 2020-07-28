from docx import Document
from json import loads

class CVWritter():
	def __init__(self, data, path=None):
		self.path =  path
		self.data = data
		self.CVmodel = Document('app/static/Model/modelo.docx')
		

	def createCV(self):
		self.escrever_dados_pessoais()
		self.escrever_formacao_academica()
		self.escrever_experiencia_profisional()
		self.escrever_projetos()
		self.escrever_habilidades()

	#Dados Pessoais
	def escrever_dados_pessoais(self):
		self.CVmodel.paragraphs[0].text = self.data['pessoais']['nome']
		self.CVmodel.paragraphs[1].text = ''.join([self.data['area de interesse'][f'{x+1}']+' | ' for x in range(len(self.data['area de interesse']))])[:-3]

		pessoais = [self.data['pessoais']['cidade_estado'], self.data['pessoais']['telefone'], 
					self.data['pessoais']['email'], self.data['pessoais']['linkedin'], self.data['pessoais']['github']]
		for i in range(2, 7):
			self.CVmodel.paragraphs[i].text += pessoais[i-2]
		
	#Formação academica
	def escrever_formacao_academica(self):
		str_formacao = ''
		if self.__validate_conteiner(self.data['formacao'], 'FORMAÇÃO'):
			for x in range(len(self.data['formacao'])):
				x = str(x+1)
				self.data_using = self.data['formacao']

				str_formacao += f"Curso: {self.data_using[x]['curso']}\n"
				str_formacao += f"{self.data_using[x]['instituicao']} - {self.data_using[x]['cidade']} - {self.data_using[x]['ano']}"
				if self.data_using[x]['horas']:
					str_formacao += f" - {self.data_using[x]['horas']} horas\n\n"
				else:
					str_formacao += '\n\n'
			self.CVmodel.paragraphs[self.__find_index('FORMAÇÃO')].text = str_formacao[:-2]


	#EXPERIÊNCIA PROFISSIONAL
	def escrever_experiencia_profisional(self):
		str_experiencia = ''
		if self.__validate_conteiner(self.data['empregos'], 'EXPERIÊNCIA'):
			for x  in range(len(self.data['empregos'])):
				x = str(x+1)
				self.data_using = self.data['empregos']
				str_experiencia += f"{self.data_using[x]['empresa']} - {self.data_using[x]['cargo']}\nInicio: {self.data_using[x]['entrada']}"
				
				if 'dias' in self.data_using[x]['saida'].split():
					str_experiencia += f' |  Até os dias atuais\n'
				else:
					str_experiencia += f" | Término: {self.data_using[x]['saida']}\n"

				str_experiencia += f"Responsabilidades: {self.data_using[x]['resumo']}\n\n"

			self.CVmodel.paragraphs[self.__find_index('EXPERIÊNCIA')].text = str_experiencia[:-2]

	#HABILIDADES
	def escrever_habilidades(self):
		str_habilidades = ''
		if self.__validate_conteiner(self.data['habilidades'], 'HABILIDADES'):
			for x in range(len(self.data['habilidades'])):
				self.data_using = self.data['habilidades']
				x = str(x+1)
				str_habilidades += f"{self.data_using[x]['tecnologia']} - {self.data_using[x]['nivel']}\n"
			self.CVmodel.paragraphs[self.__find_index('HABILIDADES')].text = str_habilidades[:-1]

	#PROJETOS DESENVOLVIDOS
	def escrever_projetos(self):
		str_projetos = ''
		if self.__validate_conteiner(self.data['projetos'], 'PROJETOS'):
			for x in range(len(self.data['projetos'])):
				self.data_using = self.data['projetos']
				x = str(x+1)

				str_projetos += f"Projeto: {self.data_using[x]['nome']}\nTecnologias: {self.data_using[x]['tecnologias']}\n"
				str_projetos += f"Resumo: {self.data_using[x]['resumo']}\nLink: {self.data_using[x]['link']}\n\n"
			self.CVmodel.paragraphs[self.__find_index('PROJETOS')].text = str_projetos[:-2]

	def save(self, filename=None):
		if not filename:
			filename = f"{self.path}/Model-{self.data['pessoais']['nome']}.docx"
		else:
			filename = f"{self.path}/Model-{filename}.docx"
			
		self.CVmodel.save(filename)
		
		return filename

	def __find_index(self, key):
		writer = False
		for i in range(len(self.CVmodel.paragraphs)):
			if key in self.CVmodel.paragraphs[i].text.split():
				return i+1

	def __validate_conteiner(self, dados, conteiner):
		if dados:
			return True
		else:
			inicio = self.__find_index(conteiner)-2
			fim = inicio + 2
			for i in range(inicio, fim+1):
				self.CVmodel.paragraphs[i].text = ''
			
			return False

	def __repr__(self):
		return f"< Curriculum - {self.data['pessoais']['nome']} >"


if __name__ == '__main__':
	data = loads(jsonfields)
	model = Document('Model/modelo.docx')
	Pedro = CVWritter(data, model)
	Pedro.createCV()
	caminho = Pedro.save()

