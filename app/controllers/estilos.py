from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx import Document
from os import path

class Padrao():
	def __init__(self, CV_filename, path=None):
		self.path = path
		self.CV_filename = CV_filename
		self.paragraphs = None
		self.document = Document()
		self.__gerarBase()

	def aplicar(self):
		self.subtitle()
		self.pessoais()
		self.escrever_conteiners()

	def subtitle(self):

		for i in range(len(self.document.paragraphs)):
			
			if i == 0:
				self.document.paragraphs[i].paragraph_format.space_before = Pt(1)
				self.document.paragraphs[i].paragraph_format.space_after = Pt(1)
				self.document.paragraphs[i].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
				run = self.document.paragraphs[i].add_run(self.paragraphs[i])
				run.font.name = 'Calibri'
				run.font.size = Pt(20)

			elif i == 1:
				self.document.paragraphs[i].paragraph_format.space_before = Pt(1)
				self.document.paragraphs[i].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
				run = self.document.paragraphs[i].add_run(self.paragraphs[i])
				run.font.name = 'Calibri'
				run.font.size = Pt(13)
				run.font.color.rgb = RGBColor(0x04, 0x89, 0xB1)

			elif self.paragraphs[i][-3:] == '___':

				run_ = self.document.paragraphs[i].add_run(self.paragraphs[i] + '______________')
				run_.font.color.rgb = RGBColor(0x41, 0x69, 0xE1)
				run_.font.name = 'Calibri'

				self.document.paragraphs[i+1].paragraph_format.left_indent = Inches(0.5)
				run = self.document.paragraphs[i+1].add_run(self.paragraphs[i+1])
				run.font.name = 'Calibri'
				run.font.size = Pt(14)
				run.font.color.rgb = RGBColor(0x41, 0x69, 0xE1)

	def pessoais(self):

		for i in range(len(self.document.paragraphs)):
			
			if i >= 2 and i <= 6:
				self.document.paragraphs[i].paragraph_format.space_before = Pt(1)
				self.document.paragraphs[i].paragraph_format.space_after = Pt(1)
				run1 = self.document.paragraphs[i].add_run(self.paragraphs[i].split(':')[0]+': ')
				run1.font.name = 'Calibri'
				run1.font.size = Pt(11)
				run1.font.bold = True

				run2 = self.document.paragraphs[i].add_run(self.paragraphs[i].split(':')[1])
				run2.font.name = 'Cambria (Corpo)'
				run2.font.size = Pt(11)

	def escrever_conteiners(self):
		writer = False
		i_pular = []
		for i in range(len(self.document.paragraphs)):
			if '___' == self.paragraphs[i][-3:]:
				writer = True
				i_pular.append(i)
				i_pular.append(i+1)

			elif '___' in self.paragraphs[i][-3:]: 
				writer = False

			if writer and i not in i_pular:
				self.document.paragraphs[i].paragraph_format.space_before = Pt(1)
				self.document.paragraphs[i].paragraph_format.space_after = Pt(1)
				run = self.document.paragraphs[i].add_run(self.paragraphs[i])
	
	def save(self):
		self.CV_filename = 'CV-'+''.join(self.CV_filename.split('-')[1])

		try:
			self.document.save(f"../Files/{self.CV_filename}")

		except Exception:
			print(f"{self.path}/{self.CV_filename}")
			self.document.save(f"{self.path}/{self.CV_filename}")

		return self.CV_filename

	def __gerarBase(self):
		try:
			CV = Document(f'../Files/{self.CV_filename}')
		except Exception:
			CV = Document(f'{self.path}/{self.CV_filename}')

		self.paragraphs = []
		controle_quebra = 0
		for paragraph in CV.paragraphs:
			for line in paragraph.text.split('\n'):
				if line == '':
					controle_quebra += 1

				else:
					controle_quebra = 0

				if controle_quebra == 3:
					self.paragraphs.pop()
					self.paragraphs.pop()

				self.paragraphs.append(line)

		for paragraph in self.paragraphs:
			self.document.add_paragraph()


if __name__ == '__main__':
	
	filename = 'Model-José Pedro da Silva Gomes.docx'

	CV = Padrao(filename)
	CV.aplicar()
	filename = CV.save()



